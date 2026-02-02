import asyncio
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from urllib.parse import urljoin, urlparse
import os

async def capture_website():
    print("Starting website documentation...")

    base_url = "https://www.yokehealth.com"
    pages_data = []

    async with async_playwright() as p:
        # Launch browser
        browser = await p.chromium.launch(headless=True)
        page = await browser.new_page()
        page.set_default_timeout(60000)  # 60 second timeout

        print(f"\n1. Capturing homepage: {base_url}")

        # Go to homepage
        await page.goto(base_url, wait_until="networkidle")
        await page.wait_for_timeout(2000)  # Wait 2 seconds for any animations

        # Get homepage title
        home_title = await page.title()
        print(f"   Title: {home_title}")

        # Take full page screenshot of homepage
        home_screenshot = "homepage.png"
        await page.screenshot(path=home_screenshot, full_page=True)
        print(f"   Screenshot saved: {home_screenshot}")

        pages_data.append({
            "title": home_title,
            "url": base_url,
            "screenshot": home_screenshot
        })

        # Get all links from homepage
        print("\n2. Finding all links on homepage...")
        links = await page.evaluate('''() => {
            const links = Array.from(document.querySelectorAll('a[href]'));
            return links.map(link => ({
                href: link.href,
                text: link.textContent.trim()
            }));
        }''')

        # Filter links - only internal links from homepage
        internal_links = []
        home_domain = urlparse(base_url).netloc

        for link in links:
            href = link['href']
            parsed = urlparse(href)

            # Only internal links, no anchors, no duplicates
            if (parsed.netloc == home_domain or not parsed.netloc) and \
               not href.startswith('#') and \
               href not in [l['url'] for l in internal_links] and \
               href != base_url and href != base_url + '/':

                full_url = urljoin(base_url, href)
                internal_links.append({
                    'url': full_url,
                    'text': link['text']
                })

        print(f"   Found {len(internal_links)} unique internal links")

        # Capture each linked page
        print("\n3. Capturing linked pages...")
        for i, link in enumerate(internal_links, 1):
            try:
                print(f"\n   [{i}/{len(internal_links)}] {link['url']}")

                await page.goto(link['url'], wait_until="networkidle")
                await page.wait_for_timeout(2000)

                # Get page title
                page_title = await page.title()
                print(f"   Title: {page_title}")

                # Take full page screenshot
                screenshot_name = f"page_{i}.png"
                await page.screenshot(path=screenshot_name, full_page=True)
                print(f"   Screenshot saved: {screenshot_name}")

                pages_data.append({
                    "title": page_title,
                    "url": link['url'],
                    "screenshot": screenshot_name
                })

            except Exception as e:
                print(f"   ERROR: Could not capture {link['url']}: {str(e)}")
                continue

        await browser.close()

    # Create Word document
    print("\n4. Creating Word document...")
    doc = Document()

    # Add title page
    title = doc.add_heading('Yoke Health Website Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {asyncio.get_event_loop().time()}")
    doc.add_paragraph(f"Total Pages: {len(pages_data)}")
    doc.add_page_break()

    # Add each page
    for i, page_data in enumerate(pages_data, 1):
        print(f"   Adding page {i}/{len(pages_data)}: {page_data['title']}")

        # Add page title
        doc.add_heading(page_data['title'], level=1)

        # Add URL
        url_para = doc.add_paragraph()
        url_para.add_run('URL: ').bold = True
        url_para.add_run(page_data['url'])

        # Add spacing
        doc.add_paragraph()

        # Add screenshot
        if os.path.exists(page_data['screenshot']):
            try:
                doc.add_picture(page_data['screenshot'], width=Inches(6.5))
            except Exception as e:
                doc.add_paragraph(f"[Screenshot error: {str(e)}]")

        # Add page break for next page
        if i < len(pages_data):
            doc.add_page_break()

    # Save document
    output_file = "YokeHealth_Website_Documentation.docx"
    doc.save(output_file)
    print(f"\n5. Document saved: {output_file}")

    # Cleanup screenshots
    print("\n6. Cleaning up temporary files...")
    for page_data in pages_data:
        if os.path.exists(page_data['screenshot']):
            os.remove(page_data['screenshot'])

    print("\n" + "="*60)
    print("WEBSITE DOCUMENTATION COMPLETE!")
    print("="*60)
    print(f"Output: {output_file}")
    print(f"Total pages documented: {len(pages_data)}")

    return output_file

if __name__ == "__main__":
    asyncio.run(capture_website())
