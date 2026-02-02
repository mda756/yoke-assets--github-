import asyncio
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import time

async def capture_website():
    print("Starting website documentation...")

    # Define pages to capture
    pages_to_capture = [
        {"name": "Home", "url": "https://yokehealth.com"},
        {"name": "AI Healthcare Platforms", "url": "https://yokehealth.com/ai-healthcare-platform-2/"},
        {"name": "Our Work", "url": "https://yokehealth.com/case-studies-digital-pharma-healthcare/"},
        {"name": "Agency", "url": "https://yokehealth.com/agency-services/"},
        {"name": "Biotech", "url": "https://yokehealth.com/biotechs/"},
        {"name": "Team", "url": "https://yokehealth.com/your-team/"},
        {"name": "Contact", "url": "https://yokehealth.com/contact-us/"},
    ]

    pages_data = []

    async with async_playwright() as p:
        # Launch browser
        browser = await p.chromium.launch(headless=True)
        page = await browser.new_page()
        page.set_default_timeout(60000)

        # Capture each page
        for i, page_info in enumerate(pages_to_capture, 1):
            try:
                print(f"\n[{i}/{len(pages_to_capture)}] Capturing: {page_info['name']}")
                print(f"   URL: {page_info['url']}")

                # Navigate to page
                await page.goto(page_info['url'], wait_until="networkidle")
                await page.wait_for_timeout(3000)  # Wait for any animations/lazy loading

                # Get page title
                page_title = await page.title()
                print(f"   Title: {page_title}")

                # Take full page screenshot
                screenshot_name = f"page_{i}_{page_info['name'].replace(' ', '_')}.png"
                await page.screenshot(path=screenshot_name, full_page=True)
                print(f"   Screenshot saved: {screenshot_name}")

                pages_data.append({
                    "title": page_title,
                    "url": page_info['url'],
                    "screenshot": screenshot_name,
                    "nav_name": page_info['name']
                })

            except Exception as e:
                print(f"   ERROR: Could not capture {page_info['url']}")
                print(f"   {str(e)[:200]}")
                continue

        await browser.close()

    # Create Word document
    print("\n\nCreating Word document...")
    doc = Document()

    # Add title page
    title = doc.add_heading('Yoke Health Website Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total Pages: {len(pages_data)}")
    doc.add_paragraph()

    # Add table of contents
    doc.add_heading('Table of Contents', level=1)
    for page_data in pages_data:
        doc.add_paragraph(f"{page_data['nav_name']} - {page_data['url']}", style='List Number')

    doc.add_page_break()

    # Add each page
    for i, page_data in enumerate(pages_data, 1):
        print(f"   Adding to document: {page_data['nav_name']}")

        # Add page heading
        doc.add_heading(f"{page_data['nav_name']}", level=1)

        # Add title
        title_para = doc.add_paragraph()
        title_para.add_run('Page Title: ').bold = True
        title_para.add_run(page_data['title'])

        # Add URL
        url_para = doc.add_paragraph()
        url_para.add_run('URL: ').bold = True
        url_para.add_run(page_data['url'])

        # Add spacing
        doc.add_paragraph()

        # Add screenshot
        if os.path.exists(page_data['screenshot']):
            try:
                doc.add_picture(page_data['screenshot'], width=Inches(6.5))
            except Exception as e:
                doc.add_paragraph(f"[Screenshot error: {str(e)}]")
        else:
            doc.add_paragraph("[Screenshot not found]")

        # Add page break for next page
        if i < len(pages_data):
            doc.add_page_break()

    # Save document
    output_file = "YokeHealth_Website_Documentation.docx"
    doc.save(output_file)
    print(f"\nDocument saved: {output_file}")

    # Cleanup screenshots
    print("\nCleaning up temporary files...")
    for page_data in pages_data:
        if os.path.exists(page_data['screenshot']):
            os.remove(page_data['screenshot'])

    print("\n" + "="*70)
    print("WEBSITE DOCUMENTATION COMPLETE!")
    print("="*70)
    print(f"\nOutput File: {output_file}")
    print(f"Total Pages: {len(pages_data)}")
    print("\nPages included:")
    for page_data in pages_data:
        print(f"  - {page_data['nav_name']}")

    return output_file

if __name__ == "__main__":
    asyncio.run(capture_website())
